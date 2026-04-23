"""Microsoft Word (.docx) extractor backed by ``python-docx``.

This extractor is optional: if ``python-docx`` is not installed the extractor
constructs but raises a clear ``RuntimeError`` from :meth:`_extract`.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from .. import EXTRACTOR_VERSIONS
from ..types import ExtractorKind
from . import Extractor, run_blocking

log = logging.getLogger(__name__)

try:
    from docx import Document  # type: ignore[import-untyped]

    _HAS_DOCX = True
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore[assignment]
    _HAS_DOCX = False


class DocxExtractor(Extractor):
    kind = ExtractorKind.DOCX
    version = EXTRACTOR_VERSIONS["docx"]

    async def _extract(
        self,
        file_path: Path,
        options: dict[str, Any],
        warnings: list[str],
    ) -> dict[str, Any]:
        if not _HAS_DOCX:
            raise RuntimeError(
                "python-docx not installed; install with "
                "`pip install datatree-multimodal[docx]`"
            )
        include_tables = bool(options.get("include_tables", True))
        return await run_blocking(
            self._extract_sync, file_path, include_tables, warnings
        )

    @staticmethod
    def _extract_sync(
        file_path: Path,
        include_tables: bool,
        warnings: list[str],
    ) -> dict[str, Any]:
        assert Document is not None
        try:
            doc = Document(str(file_path))
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"failed to open docx: {exc}") from exc

        paragraphs: list[dict[str, Any]] = []
        for p in doc.paragraphs:
            text = p.text or ""
            paragraphs.append(
                {
                    "text": text,
                    "style": getattr(p.style, "name", None),
                }
            )

        tables: list[dict[str, Any]] = []
        if include_tables:
            for t_idx, table in enumerate(doc.tables):
                rows: list[list[str]] = []
                try:
                    for row in table.rows:
                        rows.append([cell.text for cell in row.cells])
                except Exception as exc:  # noqa: BLE001
                    warnings.append(f"table {t_idx}: {exc}")
                    continue
                tables.append(
                    {
                        "index": t_idx,
                        "row_count": len(rows),
                        "col_count": (len(rows[0]) if rows else 0),
                        "rows": rows,
                    }
                )

        full_text = "\n".join(p["text"] for p in paragraphs if p["text"])
        return {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "char_count": len(full_text),
            "text": full_text,
            "paragraphs": paragraphs,
            "tables": tables,
        }


__all__ = ["DocxExtractor"]
